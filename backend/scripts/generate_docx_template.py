import docx
import os

def create_template():
    # Make sure output directory exists
    output_dir = "../frontend/public/templates"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "resume_template.docx")

    doc = docx.Document()

    # Header / Contact Info
    doc.add_heading('{{name}}', 0)
    p = doc.add_paragraph()
    p.add_run('{{email}} | {{phone}} | {{location}}\n')
    p.add_run('LinkedIn: {{linkedin}} | GitHub: {{github}}')

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('{{summary}}')

    # Experience (Loop)
    doc.add_heading('Experience', level=1)
    # Start loop
    doc.add_paragraph('{#experience}')
    # Job Title and Company
    p_job = doc.add_paragraph()
    p_job.add_run('{{title}} - {{company}}').bold = True
    p_job.add_run(' ({{date}})')
    
    # Nested bullet loop
    doc.add_paragraph('{#bullets}')
    doc.add_paragraph('{{text}}', style='List Bullet')
    doc.add_paragraph('{/bullets}')
    
    # End experience loop
    doc.add_paragraph('{/experience}')

    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph('{#education}')
    p_edu = doc.add_paragraph()
    p_edu.add_run('{{degree}} - {{institution}}').bold = True
    p_edu.add_run(', {{year}}')
    doc.add_paragraph('{/education}')

    # Skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Technical: {{technical_skills}}')
    doc.add_paragraph('Soft Skills: {{soft_skills}}')
    doc.add_paragraph('Tools: {{tools}}')

    # Save the document
    doc.save(output_path)
    print(f"Generated template at {output_path}")

if __name__ == "__main__":
    create_template()
